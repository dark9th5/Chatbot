"""
File processing utilities - Extract text from PDF and DOCX files
"""

import os
from typing import Optional
import pdfplumber
from docx import Document


def extract_text_from_pdf(file_path: str) -> Optional[str]:
    """
    Trích xuất văn bản từ file PDF
    
    Args:
        file_path: Đường dẫn đến file PDF
        
    Returns:
        str: Nội dung văn bản hoặc None nếu lỗi
    """
    try:
        text_content = []
        
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_content.append(text)
        
        full_text = '\n\n'.join(text_content)
        return full_text if full_text.strip() else None
        
    except Exception as e:
        print(f"Error extracting PDF: {e}")
        return None


def extract_text_from_docx(file_path: str) -> Optional[str]:
    """
    Trích xuất văn bản từ file DOCX
    
    Args:
        file_path: Đường dẫn đến file DOCX
        
    Returns:
        str: Nội dung văn bản hoặc None nếu lỗi
    """
    try:
        doc = Document(file_path)
        
        # Lấy text từ tất cả paragraphs
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        # Lấy text từ tables (nếu có)
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    table_texts.append(row_text)
        
        # Kết hợp tất cả
        all_text = paragraphs + table_texts
        full_text = '\n\n'.join(all_text)
        
        return full_text if full_text.strip() else None
        
    except Exception as e:
        print(f"Error extracting DOCX: {e}")
        return None


def process_uploaded_file(file_path: str, file_type: str) -> Optional[str]:
    """
    Xử lý file upload và trích xuất text
    
    Args:
        file_path: Đường dẫn đến file
        file_type: Loại file ('pdf' hoặc 'docx')
        
    Returns:
        str: Nội dung văn bản hoặc None nếu lỗi
    """
    if not os.path.exists(file_path):
        return None
    
    file_type = file_type.lower()
    
    if file_type == 'pdf':
        return extract_text_from_pdf(file_path)
    elif file_type in ['docx', 'doc']:
        return extract_text_from_docx(file_path)
    else:
        return None


def get_file_info(file_path: str) -> dict:
    """Lấy thông tin về file"""
    if not os.path.exists(file_path):
        return {}
    
    stat = os.stat(file_path)
    return {
        'size': stat.st_size,
        'size_mb': round(stat.st_size / (1024 * 1024), 2),
        'modified': stat.st_mtime
    }
